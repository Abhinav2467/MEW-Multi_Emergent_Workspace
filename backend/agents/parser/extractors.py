"""Text extraction for PDF and DOCX resumes."""

from __future__ import annotations

from pathlib import Path

import fitz
from docx import Document

from backend.agents.parser.errors import EmptyResumeError, ExtractionError, UnsupportedFileTypeError

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def detect_file_type(path: Path) -> str:
    extension = path.suffix.lower()
    if extension == ".pdf":
        return "pdf"
    if extension == ".docx":
        return "docx"
    raise UnsupportedFileTypeError(
        context={"filename": path.name, "extension": extension or "<none>"}
    )


def extract_pdf_text(path: Path) -> str:
    try:
        document = fitz.open(path)
    except Exception as exc:
        raise ExtractionError(
            "Could not open PDF file.",
            context={"path": str(path), "error": str(exc)},
        ) from exc
    try:
        pages = [page.get_text("text") for page in document]
    except Exception as exc:
        raise ExtractionError(
            "Could not extract PDF text.",
            context={"path": str(path), "error": str(exc)},
        ) from exc
    finally:
        document.close()
    return "\n".join(page.strip() for page in pages if page.strip()).strip()


def extract_docx_text(path: Path) -> str:
    try:
        document = Document(path)
    except Exception as exc:
        raise ExtractionError(
            "Could not open DOCX file.",
            context={"path": str(path), "error": str(exc)},
        ) from exc

    paragraphs = [p.text.strip() for p in document.paragraphs]
    table_cells: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    table_cells.append(text)
    lines = [line for line in paragraphs + table_cells if line]
    return "\n".join(lines).strip()


def extract_text_from_file(path: Path) -> tuple[str, str]:
    file_type = detect_file_type(path)
    text = extract_pdf_text(path) if file_type == "pdf" else extract_docx_text(path)
    if not text.strip():
        raise EmptyResumeError(context={"filename": path.name, "file_type": file_type})
    return text, file_type
