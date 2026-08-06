"""DOCX text extraction."""

from pathlib import Path

from docx import Document

from resume_parser_agent.errors import ExtractionError


def extract_docx_text(path: Path) -> str:
    """Extract paragraph and table text from a DOCX file."""

    try:
        document = Document(path)
    except Exception as exc:  # pragma: no cover - python-docx errors vary
        raise ExtractionError(
            "Could not open DOCX file.",
            context={"path": str(path), "error": str(exc)},
        ) from exc

    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    table_cells: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    table_cells.append(text)

    lines = [line for line in paragraphs + table_cells if line]
    return "\n".join(lines).strip()
